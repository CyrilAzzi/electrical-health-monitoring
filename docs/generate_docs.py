"""Génère les documents Word et PDF du projet EHM.

Lancer avec: python -X utf8 docs/generate_docs.py
"""

import os
import sys

if not sys.flags.utf8_mode:
    import subprocess
    result = subprocess.run([sys.executable, "-X", "utf8", os.path.abspath(__file__)])
    sys.exit(result.returncode)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from fpdf import FPDF

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# =============================================================================
# CONTENU PARTAGÉ
# =============================================================================

TITLE = "Electrical Health Monitoring"
SUBTITLE = "Plateforme de monitoring intelligent d'équipements électriques"
VERSION = "v1.0 — Mai 2026"


def _sections():
    """Retourne le contenu structuré de toutes les sections."""
    return [
        {
            "title": "Résumé exécutif",
            "content": [
                "EHM est une plateforme logicielle qui surveille en temps réel la santé "
                "des panneaux électriques triphasés. Des capteurs installés sur le panneau "
                "mesurent le courant et la température, puis transmettent les données sans fil "
                "vers un serveur qui analyse, détecte les anomalies et génère des alertes "
                "avant qu'une panne ne survienne.",
                "",
                "Le système calcule un score de santé de 0 à 100 pour chaque équipement, "
                "permettant aux gestionnaires de bâtiments de prioriser la maintenance "
                "et d'éviter les pannes coûteuses.",
            ],
        },
        {
            "title": "Le problème",
            "content": [
                "Les panneaux électriques des bâtiments commerciaux et institutionnels "
                "vieillissent silencieusement. Les signes avant-coureurs d'une panne "
                "(surchauffe, surcharge, déséquilibre) existent souvent des heures ou "
                "des jours avant l'incident, mais personne ne les voit.",
                "",
                "Conséquences :",
                "• Arrêt non planifié des opérations (coût : milliers de $/heure en usine)",
                "• Risque d'incendie électrique",
                "• Réclamations d'assurance refusées par manque de suivi",
                "• Interventions d'urgence 3x plus chères que la maintenance préventive",
            ],
        },
        {
            "title": "Notre solution",
            "content": [
                "Un kit de capteurs à ~100 $ par panneau, connecté à une plateforme "
                "logicielle intelligente qui fonctionne 24/7.",
                "",
                "Ce que le système surveille :",
                "• Courant sur les 3 phases (A, B, C) — détecte surcharge et déséquilibre",
                "• Température à 3 points dans le panneau — détecte surchauffe et tendance",
                "• Tension sur les phases (optionnel, phase 2)",
                "• Batterie de secours (optionnel, si UPS présent)",
                "",
                "Ce que le système fait automatiquement :",
                "• Analyse chaque mesure en temps réel",
                "• Déclenche des alertes selon 6 règles configurables",
                "• Calcule un score de santé de 0 à 100",
                "• Affiche un tableau de bord visuel (Grafana)",
                "• Résout les alertes automatiquement quand le problème disparaît",
            ],
        },
        {
            "title": "Architecture technique",
            "content": [
                "Le système est composé de 4 couches :",
                "",
                "1. CAPTEURS (sur le panneau)",
                "   ESP32 + 3 capteurs de courant (CT) + 3 sondes de température",
                "   → Envoie les données toutes les 3 secondes via WiFi / MQTT",
                "",
                "2. BROKER MQTT (serveur)",
                "   Eclipse Mosquitto — reçoit et redistribue les messages",
                "",
                "3. BACKEND (serveur)",
                "   FastAPI (Python) — stocke les mesures, évalue les règles d'alerte,",
                "   calcule le score de santé, expose l'API REST",
                "",
                "4. BASE DE DONNÉES + DASHBOARD",
                "   PostgreSQL/TimescaleDB — stockage optimisé pour les séries temporelles",
                "   Grafana — tableaux de bord visuels avec courbes et alertes",
            ],
        },
        {
            "title": "Les 6 règles d'alerte",
            "content": [
                "Chaque règle a un seuil configurable par équipement :",
                "",
                "1. SURCOURANT — Courant > 80% du nominal",
                "   → Détecte une surcharge avant que le disjoncteur ne saute",
                "",
                "2. DÉSÉQUILIBRE DE COURANT — Écart entre phases > 10%",
                "   → Signe de charge mal répartie ou de problème de connexion",
                "",
                "3. TEMPÉRATURE ÉLEVÉE — Capteur > 60°C",
                "   → Alerte surchauffe directe, risque d'incendie",
                "",
                "4. TENDANCE DE TEMPÉRATURE — Hausse continue sur 4+ mesures",
                "   → Détecte une dégradation progressive avant d'atteindre le seuil",
                "",
                "5. BATTERIE FAIBLE — Tension < 12.2V (si UPS présent)",
                "   → L'onduleur ne protégera pas en cas de coupure",
                "",
                "6. TENSION ANORMALE — Écart > 10% du nominal (si capteurs installés)",
                "   → Problème d'alimentation en amont",
            ],
        },
        {
            "title": "Score de santé (0-100)",
            "content": [
                "Le score combine 4 ou 5 sous-scores selon les capteurs installés :",
                "",
                "Avec CTs + température seulement (MVP) :",
                "• Courant : 31% — pénalité si proche du nominal",
                "• Équilibre des phases : 25% — pénalité si déséquilibré",
                "• Température : 31% — pénalité si élevée",
                "• Tendance température : 13% — pénalité si hausse continue",
                "",
                "Avec tous les capteurs :",
                "• Courant : 25% | Équilibre : 20% | Température : 25%",
                "• Tendance : 10% | Batterie : 20%",
                "",
                "Interprétation :",
                "• 100 = Excellent — aucune intervention nécessaire",
                "• 70-99 = Normal — fonctionnement dans les normes",
                "• 40-69 = À surveiller — planifier une inspection",
                "• 0-39 = Critique — intervention urgente recommandée",
            ],
        },
        {
            "title": "Choix de design",
            "content": [
                "POURQUOI PAS DE MESURE DE TENSION AU MVP ?",
                "Mesurer la tension dans un panneau 347/600V nécessite des capteurs "
                "certifiés CSA/UL et une isolation galvanique rigoureuse. Les CTs (courant) "
                "sont non-invasifs et sécuritaires. En commençant sans tension, on élimine "
                "80% du risque technique et réglementaire tout en couvrant la majorité "
                "des anomalies détectables.",
                "",
                "POURQUOI MQTT ?",
                "Protocole standard IoT, léger, fiable, supporte les déconnexions "
                "temporaires. Tous les ESP32 le supportent nativement.",
                "",
                "POURQUOI DES SEUILS CONFIGURABLES ?",
                "Un panneau de data center (tolérance zéro) n'a pas les mêmes seuils "
                "qu'un panneau d'entrepôt. Chaque client peut ajuster selon ses besoins.",
                "",
                "POURQUOI LA DÉDUPLICATION ?",
                "Sans déduplication, une température à 65°C génère une alerte à chaque "
                "mesure (toutes les 3 secondes). Avec déduplication, une seule alerte "
                "est créée et reste active jusqu'à résolution.",
            ],
        },
        {
            "title": "Matériel nécessaire par panneau",
            "content": [
                "COMPOSANT                  | MODÈLE              | QTÉ | PRIX    | OÙ ACHETER",
                "Microcontrôleur            | ESP32 DevKit V1     |  1  | ~10 $   | Amazon.ca, Digikey.ca",
                "Capteurs de courant        | YHDC SCT-013-000    |  3  | ~12 $/u | Amazon.ca, AliExpress",
                "Sondes de température      | DS18B20 étanche     |  3  | ~4 $/u  | Amazon.ca, AliExpress",
                "Résistances pour CTs       | 33 ohms             |  3  | ~1 $ lot| Digikey.ca",
                "Alimentation               | 5V USB              |  1  | ~8 $    | Amazon.ca",
                "Prototypage                | Breadboard + fils   |  1  | ~10 $   | Amazon.ca",
                "Boîtier                    | ABS projet          |  1  | ~8 $    | Amazon.ca",
                "",
                "TOTAL PAR PANNEAU : ~100 à 130 $ CAD",
                "",
                "Pour la phase 2 (avec tension) :",
                "Transducteurs de tension industriels (CR Magnetics) : +150-200 $/panneau",
                "Certification CSA nécessaire avant installation commerciale.",
            ],
        },
        {
            "title": "Installation physique",
            "content": [
                "OÙ INSTALLER LES CAPTEURS :",
                "",
                "• CT Phase A : clipsé autour du conducteur noir (phase A) à l'entrée du panneau",
                "• CT Phase B : clipsé autour du conducteur rouge (phase B)",
                "• CT Phase C : clipsé autour du conducteur bleu (phase C)",
                "• Température 1 : collée sur la barre bus principale (point le plus chaud)",
                "• Température 2 : collée sur le disjoncteur principal",
                "• Température 3 : suspendue dans le panneau (température ambiante)",
                "• ESP32 : dans un boîtier ABS, à l'intérieur ou à côté du panneau",
                "",
                "IMPORTANT :",
                "• L'ouverture d'un panneau électrique nécessite un électricien licencié au Québec",
                "• Les CTs se clipsent autour du fil SANS le couper — installation non-invasive",
                "• L'ESP32 se connecte au WiFi du bâtiment — aucun câblage réseau",
            ],
        },
        {
            "title": "Clients cibles",
            "content": [
                "IDÉAL (panneaux triphasés, coût de panne élevé) :",
                "• Immeubles commerciaux — tours de bureaux, centres commerciaux",
                "• Institutions — écoles, hôpitaux, CHSLD, universités",
                "• Industriel — usines, ateliers, entrepôts automatisés",
                "• Data centers — tolérance zéro aux interruptions",
                "• Multi-logements — condos 50+ unités, gestion d'immeubles",
                "",
                "MOINS ADAPTÉ :",
                "• Maisons unifamiliales — généralement monophasé, coût non justifié",
                "• Petits commerces — souvent monophasé, budget limité",
            ],
        },
        {
            "title": "Modèle d'affaires recommandé",
            "content": [
                "OPTION RECOMMANDÉE : Matériel + Abonnement SaaS",
                "",
                "• Kit matériel : 500 à 1 500 $ (installation incluse par un électricien partenaire)",
                "• Abonnement mensuel : 50 à 500 $/mois selon la taille du site",
                "• Le logiciel (alertes, score, dashboard) justifie l'abonnement récurrent",
                "",
                "Pourquoi pas juste vendre le kit ?",
                "→ Marges faibles sur le matériel, pas de revenus récurrents",
                "→ Le SaaS crée une relation long terme avec le client",
                "→ Les mises à jour logicielles (nouvelles règles, IA) ajoutent de la valeur",
                "",
                "Argument de vente clé :",
                "\"Une intervention d'urgence coûte 3x plus cher qu'une maintenance planifiée.\"",
                "\"Notre système vous prévient avant la panne pour ~2 $/jour.\"",
            ],
        },
        {
            "title": "Prochaines étapes",
            "content": [
                "COURT TERME (1-2 mois) :",
                "• Développer le firmware ESP32 (Arduino/ESP-IDF)",
                "• Acheter 1 kit de test (~100 $) et valider sur un panneau réel",
                "• Trouver un électricien partenaire pour les installations",
                "",
                "MOYEN TERME (3-6 mois) :",
                "• Ajouter l'authentification (JWT) et le chiffrement (TLS)",
                "• Notifications par email et SMS (Twilio, SendGrid)",
                "• Interface web simplifiée pour les clients (pas Grafana brut)",
                "• Premier client pilote gratuit pour valider le produit",
                "",
                "LONG TERME (6-12 mois) :",
                "• Déploiement cloud (AWS/Azure) pour le multi-client",
                "• Application mobile avec notifications push",
                "• Ajout de la mesure de tension (capteurs industriels certifiés)",
                "• Certifications CSA/UL pour le matériel",
                "• Intégration IA pour la prédiction de pannes",
            ],
        },
    ]


# =============================================================================
# GÉNÉRATION WORD
# =============================================================================

def generate_word():
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Page titre
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SUBTITLE)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(80, 80, 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(VERSION)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # Table des matières (manuelle)
    doc.add_heading("Table des matières", level=1)
    for i, section in enumerate(_sections(), 1):
        p = doc.add_paragraph(f"{i}. {section['title']}")
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # Sections
    for section in _sections():
        doc.add_heading(section["title"], level=1)
        for line in section["content"]:
            if line == "":
                doc.add_paragraph()
            elif line.startswith("•"):
                p = doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("→"):
                p = doc.add_paragraph(line[2:], style="List Bullet 2")
            elif "|" in line and "COMPOSANT" in line:
                # En-tête du tableau matériel
                _add_material_table(doc)
                break
            elif line.isupper() or (line.endswith(":") and len(line) < 60):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(11)
            else:
                doc.add_paragraph(line)

    path = os.path.join(OUTPUT_DIR, "EHM_Guide_Projet.docx")
    doc.save(path)
    print(f"Word généré : {path}")
    return path


def _add_material_table(doc):
    """Ajoute le tableau du matériel dans le document Word."""
    headers = ["Composant", "Modèle", "Qté", "Prix", "Où acheter"]
    rows = [
        ["Microcontrôleur", "ESP32 DevKit V1", "1", "~10 $", "Amazon.ca, Digikey.ca"],
        ["Capteurs de courant", "YHDC SCT-013-000", "3", "~12 $/u", "Amazon.ca, AliExpress"],
        ["Sondes température", "DS18B20 étanche", "3", "~4 $/u", "Amazon.ca, AliExpress"],
        ["Résistances CTs", "33 ohms", "3", "~1 $ lot", "Digikey.ca"],
        ["Alimentation", "5V USB", "1", "~8 $", "Amazon.ca"],
        ["Prototypage", "Breadboard + fils", "1", "~10 $", "Amazon.ca"],
        ["Boîtier", "ABS projet", "1", "~8 $", "Amazon.ca"],
    ]

    table = doc.add_table(rows=len(rows) + 1, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("TOTAL PAR PANNEAU : ~100 à 130 $ CAD")
    run.bold = True


# =============================================================================
# GÉNÉRATION PDF
# =============================================================================

FONT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")


class EHM_PDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_font("DejaVu", "", os.path.join(FONT_DIR, "DejaVuSans.ttf"))
        self.add_font("DejaVu", "B", os.path.join(FONT_DIR, "DejaVuSans-Bold.ttf"))
        self.add_font("DejaVu", "I", os.path.join(FONT_DIR, "DejaVuSans-Oblique.ttf"))

    def _font(self, style="", size=10):
        self.set_font("DejaVu", style, size)

    def header(self):
        if self.page_no() > 1:
            self._font("I", 8)
            self.set_text_color(120, 120, 120)
            self.cell(0, 10, f"{TITLE} — {VERSION}", align="L")
            self.cell(0, 10, f"Page {self.page_no()}", align="R")
            self.ln(15)

    def footer(self):
        pass

    def section_title(self, title):
        self.set_x(self.l_margin)
        self._font("B", 16)
        self.set_text_color(0, 51, 102)
        self.cell(self.w - self.l_margin - self.r_margin, 12, title, new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(0, 51, 102)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(6)

    def _write_block(self, text):
        """multi_cell avec reset du curseur X à la marge gauche."""
        self.set_x(self.l_margin)
        self.multi_cell(w=self.w - self.l_margin - self.r_margin, h=5.5, text=text)

    def body_text(self, text):
        self._font("", 10)
        self.set_text_color(40, 40, 40)
        self._write_block(text)
        self.ln(2)

    def bold_text(self, text):
        self._font("B", 10)
        self.set_text_color(40, 40, 40)
        self._write_block(text)
        self.ln(1)

    def bullet(self, text):
        self._font("", 10)
        self.set_text_color(40, 40, 40)
        self._write_block(f"   -  {text}")

    def sub_bullet(self, text):
        self._font("", 10)
        self.set_text_color(80, 80, 80)
        self._write_block(f"        >  {text}")


def generate_pdf():
    pdf = EHM_PDF()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Page titre
    pdf.add_page()
    pdf.ln(80)
    pdf._font("B", 32)
    pdf.set_text_color(0, 51, 102)
    pdf.cell(0, 15, TITLE, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)
    pdf._font("", 14)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 10, SUBTITLE, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)
    pdf._font("I", 11)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 10, VERSION, align="C", new_x="LMARGIN", new_y="NEXT")

    # Sections
    for section in _sections():
        pdf.add_page()
        pdf.section_title(section["title"])

        for line in section["content"]:
            if line == "":
                pdf.ln(3)
            elif line.startswith(chr(8226)) or line.startswith("•"):
                pdf.bullet(line.lstrip("•").lstrip())
            elif line.startswith("→"):
                pdf.sub_bullet(line.lstrip("→").lstrip())
            elif "|" in line and "COMPOSANT" in line:
                _add_pdf_material_table(pdf)
                break
            elif line.isupper() or (line.endswith(":") and len(line) < 60):
                pdf.bold_text(line)
            else:
                pdf.body_text(line)

    path = os.path.join(OUTPUT_DIR, "EHM_Guide_Projet.pdf")
    pdf.output(path)
    print(f"PDF généré : {path}")
    return path


def _add_pdf_material_table(pdf):
    """Ajoute le tableau matériel dans le PDF."""
    headers = ["Composant", "Modèle", "Qté", "Prix", "Où acheter"]
    rows = [
        ["Microcontrôleur", "ESP32 DevKit V1", "1", "~10 $", "Amazon.ca, Digikey"],
        ["Capteurs courant", "YHDC SCT-013-000", "3", "~12 $/u", "Amazon.ca"],
        ["Sondes temp.", "DS18B20 étanche", "3", "~4 $/u", "Amazon.ca"],
        ["Résistances", "33 ohms", "3", "~1 $ lot", "Digikey.ca"],
        ["Alimentation", "5V USB", "1", "~8 $", "Amazon.ca"],
        ["Prototypage", "Breadboard + fils", "1", "~10 $", "Amazon.ca"],
        ["Boîtier", "ABS projet", "1", "~8 $", "Amazon.ca"],
    ]

    col_widths = [38, 42, 12, 22, 38]

    # Header
    pdf._font("B", 9)
    pdf.set_fill_color(0, 51, 102)
    pdf.set_text_color(255, 255, 255)
    for i, h in enumerate(headers):
        pdf.cell(col_widths[i], 8, h, border=1, fill=True, align="C")
    pdf.ln()

    # Rows
    pdf._font("", 9)
    pdf.set_text_color(40, 40, 40)
    for r, row_data in enumerate(rows):
        fill = r % 2 == 0
        if fill:
            pdf.set_fill_color(240, 245, 250)
        for i, val in enumerate(row_data):
            pdf.cell(col_widths[i], 7, val, border=1, fill=fill, align="C" if i == 2 else "L")
        pdf.ln()

    pdf.ln(5)
    pdf._font("B", 11)
    pdf.set_text_color(0, 51, 102)
    pdf.cell(0, 8, "TOTAL PAR PANNEAU : ~100 à 130 $ CAD", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # Texte phase 2
    pdf._font("", 10)
    pdf.set_text_color(40, 40, 40)
    pdf.body_text(
        "Pour la phase 2 (avec tension) : transducteurs de tension industriels "
        "(CR Magnetics) : +150-200 $/panneau. Certification CSA nécessaire."
    )


# =============================================================================
# MAIN
# =============================================================================

if __name__ == "__main__":
    print("Génération des documents EHM...")
    generate_word()
    generate_pdf()
    print("Terminé.")
